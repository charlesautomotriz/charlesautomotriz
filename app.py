import os
import re
import shutil
import sqlite3
import subprocess
import tempfile
import base64
from pathlib import Path
from datetime import datetime, date
from zoneinfo import ZoneInfo

import pandas as pd
import streamlit as st
import plotly.express as px
from docx import Document
from openpyxl import load_workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

# =========================================================
# CONFIGURACIÓN GENERAL
# =========================================================
st.set_page_config(
    page_title="Cotizaciones Charles Servicio Automotriz",
    page_icon="🔧",
    layout="wide"
)

BASE_DIR = Path(__file__).parent
DB_FILE = BASE_DIR / "cotizaciones_charles.db"
EXCEL_FILE = BASE_DIR / "base_cotizaciones_charles.xlsx"
OUTPUT_DIR = BASE_DIR / "cotizaciones_generadas"
OUTPUT_DIR.mkdir(exist_ok=True)

TEMPLATE_FILE = BASE_DIR / "Cotizacion Charles Servicio Automotriz.docx"
LOGO_HEADER = BASE_DIR / "marca charles.png"
LOGO_WATERMARK = BASE_DIR / "logo charles blanco 21 jun 2025, 23_06_15.png"

ESTADOS = ["Pendiente", "Aceptada", "Rechazada", "En revisión", "Vencida"]

# =========================================================
# ESTILO VISUAL
# =========================================================
def image_to_base64(path: Path) -> str:
    if not path.exists():
        return ""
    with open(path, "rb") as f:
        return base64.b64encode(f.read()).decode("utf-8")


def aplicar_estilos():
    watermark_b64 = image_to_base64(LOGO_WATERMARK)
    bg_css = ""
    if watermark_b64:
        bg_css = f"""
        .stApp::after {{
            content: "";
            position: fixed;
            top: 58%;
            left: 50%;
            transform: translate(-50%, -50%);
            width: 560px;
            height: 560px;
            background-image: url("data:image/png;base64,{watermark_b64}");
            background-repeat: no-repeat;
            background-position: center;
            background-size: contain;
            opacity: 0.035;
            z-index: 0;
            pointer-events: none;
        }}
        """

    st.markdown(
        f"""
        <style>
            .stApp {{ background-color: #ffffff; }}
            {bg_css}
            .main-title {{
                font-size: 34px;
                font-weight: 800;
                color: #111111;
                margin-bottom: 0px;
            }}
            .sub-title {{
                font-size: 15px;
                color: #666666;
                margin-top: 0px;
            }}
            .card-box {{
                border: 1px solid #eeeeee;
                border-radius: 14px;
                padding: 16px;
                background-color: rgba(255,255,255,0.92);
                box-shadow: 0 2px 8px rgba(0,0,0,0.04);
            }}
            .total-box {{
                border-radius: 14px;
                padding: 18px;
                background: #fff7ed;
                border-left: 7px solid #d97706;
            }}
            .small-muted {{ color: #777777; font-size: 13px; }}
        </style>
        """,
        unsafe_allow_html=True,
    )


aplicar_estilos()

# =========================================================
# UTILIDADES
# =========================================================
def ahora_santiago() -> datetime:
    return datetime.now(ZoneInfo("America/Santiago"))


def clp(valor) -> str:
    try:
        return f"$ {float(valor):,.0f}".replace(",", ".")
    except Exception:
        return "$ 0"


def limpiar_nombre_archivo(texto: str) -> str:
    texto = str(texto).strip().replace(" ", "_")
    texto = re.sub(r"[^A-Za-z0-9_\-]", "", texto)
    return texto[:80] if texto else "sin_nombre"


def numero_siguiente() -> str:
    conn = sqlite3.connect(DB_FILE)
    cur = conn.cursor()
    cur.execute("SELECT MAX(id) FROM cotizaciones")
    ultimo = cur.fetchone()[0]
    conn.close()
    siguiente = 1 if ultimo is None else int(ultimo) + 1
    return f"CSA-{siguiente:04d}"


def limpiar_df_items(df: pd.DataFrame, categoria: str) -> pd.DataFrame:
    if df is None or df.empty:
        return pd.DataFrame(columns=["categoria", "descripcion", "cantidad", "valor_unitario", "total"])

    df = df.copy()
    for col in ["descripcion", "cantidad", "valor_unitario"]:
        if col not in df.columns:
            df[col] = "" if col == "descripcion" else 0

    df["descripcion"] = df["descripcion"].fillna("").astype(str).str.strip()
    df["cantidad"] = pd.to_numeric(df["cantidad"], errors="coerce").fillna(0)
    df["valor_unitario"] = pd.to_numeric(df["valor_unitario"], errors="coerce").fillna(0)
    df = df[(df["descripcion"] != "") & (df["cantidad"] > 0)]
    df["total"] = df["cantidad"] * df["valor_unitario"]
    df["categoria"] = categoria
    return df[["categoria", "descripcion", "cantidad", "valor_unitario", "total"]]


def editor_items(titulo: str, key: str, ejemplo: str) -> pd.DataFrame:
    st.markdown(f"#### {titulo}")
    df_base = pd.DataFrame([
        {"descripcion": ejemplo, "cantidad": 1, "valor_unitario": 0},
        {"descripcion": "", "cantidad": 1, "valor_unitario": 0},
        {"descripcion": "", "cantidad": 1, "valor_unitario": 0},
    ])
    df = st.data_editor(
        df_base,
        key=key,
        num_rows="dynamic",
        use_container_width=True,
        hide_index=True,
        column_config={
            "descripcion": st.column_config.TextColumn("Descripción", width="large"),
            "cantidad": st.column_config.NumberColumn("Cantidad", min_value=0.0, step=1.0),
            "valor_unitario": st.column_config.NumberColumn("Valor unitario $", min_value=0, step=1000, format="$ %d"),
        },
    )
    return df

# =========================================================
# BASE DE DATOS Y EXCEL
# =========================================================
def get_conn():
    return sqlite3.connect(DB_FILE, check_same_thread=False)


def init_db():
    conn = get_conn()
    cur = conn.cursor()
    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS cotizaciones (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            numero_cotizacion TEXT UNIQUE NOT NULL,
            fecha TEXT NOT NULL,
            cliente TEXT NOT NULL,
            contacto TEXT,
            mail TEXT,
            patente TEXT,
            marca TEXT,
            modelo TEXT,
            anio TEXT,
            vin TEXT,
            kilometraje TEXT,
            observaciones TEXT,
            total_mano_obra REAL DEFAULT 0,
            total_repuestos REAL DEFAULT 0,
            total_otros REAL DEFAULT 0,
            total_final REAL DEFAULT 0,
            estado TEXT DEFAULT 'Pendiente',
            motivo_estado TEXT,
            fecha_estado TEXT,
            creado_en TEXT NOT NULL,
            actualizado_en TEXT,
            docx_path TEXT,
            pdf_path TEXT
        )
        """
    )
    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS items_cotizacion (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            cotizacion_id INTEGER NOT NULL,
            categoria TEXT NOT NULL,
            descripcion TEXT NOT NULL,
            cantidad REAL DEFAULT 1,
            valor_unitario REAL DEFAULT 0,
            total REAL DEFAULT 0,
            FOREIGN KEY (cotizacion_id) REFERENCES cotizaciones(id)
        )
        """
    )
    conn.commit()
    conn.close()


def guardar_cotizacion(datos: dict, items: pd.DataFrame) -> int:
    conn = get_conn()
    cur = conn.cursor()
    cur.execute(
        """
        INSERT INTO cotizaciones (
            numero_cotizacion, fecha, cliente, contacto, mail,
            patente, marca, modelo, anio, vin, kilometraje, observaciones,
            total_mano_obra, total_repuestos, total_otros, total_final,
            estado, creado_en, actualizado_en, docx_path, pdf_path
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            datos["numero_cotizacion"], datos["fecha"], datos["cliente"], datos["contacto"], datos["mail"],
            datos["patente"], datos["marca"], datos["modelo"], datos["anio"], datos["vin"], datos["kilometraje"],
            datos["observaciones"], datos["total_mano_obra"], datos["total_repuestos"], datos["total_otros"],
            datos["total_final"], "Pendiente", datos["creado_en"], datos["creado_en"], datos.get("docx_path", ""), datos.get("pdf_path", "")
        ),
    )
    cotizacion_id = cur.lastrowid

    for _, row in items.iterrows():
        cur.execute(
            """
            INSERT INTO items_cotizacion (
                cotizacion_id, categoria, descripcion, cantidad, valor_unitario, total
            ) VALUES (?, ?, ?, ?, ?, ?)
            """,
            (
                cotizacion_id,
                row["categoria"],
                row["descripcion"],
                float(row["cantidad"]),
                float(row["valor_unitario"]),
                float(row["total"]),
            ),
        )

    conn.commit()
    conn.close()
    sincronizar_excel()
    return cotizacion_id


def actualizar_archivos_cotizacion(cotizacion_id: int, docx_path: str = "", pdf_path: str = ""):
    conn = get_conn()
    cur = conn.cursor()
    cur.execute(
        "UPDATE cotizaciones SET docx_path=?, pdf_path=?, actualizado_en=? WHERE id=?",
        (docx_path, pdf_path, ahora_santiago().strftime("%Y-%m-%d %H:%M:%S"), cotizacion_id),
    )
    conn.commit()
    conn.close()
    sincronizar_excel()


def actualizar_estado(cotizacion_id: int, estado: str, motivo: str):
    conn = get_conn()
    cur = conn.cursor()
    ahora = ahora_santiago().strftime("%Y-%m-%d %H:%M:%S")
    cur.execute(
        """
        UPDATE cotizaciones
        SET estado=?, motivo_estado=?, fecha_estado=?, actualizado_en=?
        WHERE id=?
        """,
        (estado, motivo, ahora, ahora, cotizacion_id),
    )
    conn.commit()
    conn.close()
    sincronizar_excel()


def cargar_cotizaciones() -> pd.DataFrame:
    conn = get_conn()
    df = pd.read_sql_query("SELECT * FROM cotizaciones ORDER BY id DESC", conn)
    conn.close()
    return df


def cargar_items(cotizacion_id: int | None = None) -> pd.DataFrame:
    conn = get_conn()
    if cotizacion_id is None:
        df = pd.read_sql_query("SELECT * FROM items_cotizacion ORDER BY cotizacion_id DESC, id ASC", conn)
    else:
        df = pd.read_sql_query(
            "SELECT * FROM items_cotizacion WHERE cotizacion_id=? ORDER BY id ASC",
            conn,
            params=(cotizacion_id,),
        )
    conn.close()
    return df


def sincronizar_excel():
    try:
        df_cot = cargar_cotizaciones_sin_excel()
        df_items = cargar_items_sin_excel()
        with pd.ExcelWriter(EXCEL_FILE, engine="openpyxl") as writer:
            df_cot.to_excel(writer, sheet_name="cotizaciones", index=False)
            df_items.to_excel(writer, sheet_name="detalle_items", index=False)

        wb = load_workbook(EXCEL_FILE)
        header_fill = PatternFill("solid", fgColor="D97706")
        header_font = Font(color="FFFFFF", bold=True)
        thin = Side(style="thin", color="D9D9D9")
        border = Border(left=thin, right=thin, top=thin, bottom=thin)

        for ws in wb.worksheets:
            ws.freeze_panes = "A2"
            for cell in ws[1]:
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = Alignment(horizontal="center", vertical="center")
                cell.border = border

            for row in ws.iter_rows(min_row=2):
                for cell in row:
                    cell.border = border
                    cell.alignment = Alignment(vertical="center")

            for col_idx, col in enumerate(ws.columns, 1):
                max_len = 10
                for cell in col:
                    max_len = max(max_len, len(str(cell.value)) if cell.value is not None else 0)
                ws.column_dimensions[get_column_letter(col_idx)].width = min(max_len + 2, 42)

        wb.save(EXCEL_FILE)
    except Exception as e:
        st.warning(f"No fue posible actualizar la planilla Excel: {e}")


def cargar_cotizaciones_sin_excel() -> pd.DataFrame:
    conn = sqlite3.connect(DB_FILE, check_same_thread=False)
    df = pd.read_sql_query("SELECT * FROM cotizaciones ORDER BY id DESC", conn)
    conn.close()
    return df


def cargar_items_sin_excel() -> pd.DataFrame:
    conn = sqlite3.connect(DB_FILE, check_same_thread=False)
    df = pd.read_sql_query("SELECT * FROM items_cotizacion ORDER BY cotizacion_id DESC, id ASC", conn)
    conn.close()
    return df

# =========================================================
# WORD / PDF
# =========================================================
def reemplazar_texto_docx(doc: Document, reemplazos: list[tuple[str, str]]):
    def reemplazar_en_parrafo(paragraph):
        if not paragraph.runs:
            return
        texto = "".join(run.text for run in paragraph.runs)
        texto_nuevo = texto
        for buscar, reemplazar in reemplazos:
            texto_nuevo = texto_nuevo.replace(buscar, str(reemplazar))
        if texto_nuevo != texto:
            for run in paragraph.runs:
                run.text = ""
            paragraph.runs[0].text = texto_nuevo

    def procesar_tabla(table):
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    reemplazar_en_parrafo(p)
                for subtable in cell.tables:
                    procesar_tabla(subtable)

    for p in doc.paragraphs:
        reemplazar_en_parrafo(p)

    for table in doc.tables:
        procesar_tabla(table)

    for section in doc.sections:
        for p in section.header.paragraphs:
            reemplazar_en_parrafo(p)
        for table in section.header.tables:
            procesar_tabla(table)
        for p in section.footer.paragraphs:
            reemplazar_en_parrafo(p)
        for table in section.footer.tables:
            procesar_tabla(table)


def valor_item(items: pd.DataFrame, categoria: str, indice: int, campo: str) -> str:
    subset = items[items["categoria"] == categoria].reset_index(drop=True)
    if indice >= len(subset):
        return ""
    row = subset.iloc[indice]
    if campo == "descripcion":
        return str(row["descripcion"])
    if campo == "total":
        return clp(row["total"])
    return ""


def texto_items_extra(items: pd.DataFrame) -> str:
    lineas = []

    for categoria in ["Mano de obra", "Repuestos", "Otros"]:
        subset = items[items["categoria"] == categoria].reset_index(drop=True)
        if categoria in ["Mano de obra", "Repuestos"]:
            subset_extra = subset.iloc[3:] if len(subset) > 3 else pd.DataFrame()
        else:
            subset_extra = subset

        if not subset_extra.empty:
            lineas.append(f"{categoria} adicionales:")
            for _, row in subset_extra.iterrows():
                lineas.append(f"- {row['descripcion']} ({clp(row['total'])})")

    return "\n".join(lineas)


def generar_docx_cotizacion(datos: dict, items: pd.DataFrame) -> Path:
    if not TEMPLATE_FILE.exists():
        raise FileNotFoundError(f"No se encontró la plantilla: {TEMPLATE_FILE.name}")

    doc = Document(str(TEMPLATE_FILE))
    numero = datos["numero_cotizacion"]
    observaciones_base = datos.get("observaciones", "").strip()

    detalle_vehiculo = []
    if datos.get("patente"):
        detalle_vehiculo.append(f"Patente: {datos['patente']}")
    if datos.get("marca") or datos.get("modelo"):
        detalle_vehiculo.append(f"Vehículo: {datos.get('marca', '')} {datos.get('modelo', '')}".strip())
    if datos.get("anio"):
        detalle_vehiculo.append(f"Año: {datos['anio']}")
    if datos.get("kilometraje"):
        detalle_vehiculo.append(f"Kilometraje: {datos['kilometraje']}")
    if datos.get("vin"):
        detalle_vehiculo.append(f"VIN: {datos['vin']}")

    extras = texto_items_extra(items)
    obs_final = "\n".join([x for x in [observaciones_base, " | ".join(detalle_vehiculo), extras] if x])

    reemplazos = [
        ("00<<Numero de cotización>>", numero),
        ("00<<Numero de cotizacion>>", numero),
        ("<<Numero de cotización>>", numero),
        ("<<Numero de cotizacion>>", numero),
        ("<<Fecha>>", datos["fecha_mostrar"]),
        ("<<fecha>>", datos["fecha_mostrar"]),
        ("<<Nombre del cliente>>", datos["cliente"]),
        ("<<Numero de contacto>>", datos.get("contacto", "")),
        ("<<Número de contacto>>", datos.get("contacto", "")),
        ("<<Mail del cliente>>", datos.get("mail", "")),
        ("<<Trabajos a Realizar 1>>", valor_item(items, "Mano de obra", 0, "descripcion")),
        ("<<Trabajos a Realizar 2>>", valor_item(items, "Mano de obra", 1, "descripcion")),
        ("<<Trabajos a Realizar 3>>", valor_item(items, "Mano de obra", 2, "descripcion")),
        ("<<Valor total servicio 1>>", valor_item(items, "Mano de obra", 0, "total")),
        ("<<Valor total servicio 2>>", valor_item(items, "Mano de obra", 1, "total")),
        ("<<Valor total servicio 3>>", valor_item(items, "Mano de obra", 2, "total")),
        ("<<Repuestos comprometidos 1>>", valor_item(items, "Repuestos", 0, "descripcion")),
        ("<<Repuestos comprometidos 2>>", valor_item(items, "Repuestos", 1, "descripcion")),
        ("<<Repuestos comprometidos 3>>", valor_item(items, "Repuestos", 2, "descripcion")),
        ("<<Repuestos comprometido 1>>", valor_item(items, "Repuestos", 0, "descripcion")),
        ("<<Repuestos comprometido 2>>", valor_item(items, "Repuestos", 1, "descripcion")),
        ("<<Repuestos comprometido 3>>", valor_item(items, "Repuestos", 2, "descripcion")),
        ("<<Valor total repuestos comprometidos 1>>", valor_item(items, "Repuestos", 0, "total")),
        ("<<Valor total repuestos comprometidos 2>>", valor_item(items, "Repuestos", 1, "total")),
        ("<<Valor total repuestos comprometidos 3>>", valor_item(items, "Repuestos", 2, "total")),
        ("<<Valor total repuesto comprometido 1>>", valor_item(items, "Repuestos", 0, "total")),
        ("<<Valor total repuesto comprometido 2>>", valor_item(items, "Repuestos", 1, "total")),
        ("<<Valor total repuesto comprometido 3>>", valor_item(items, "Repuestos", 2, "total")),
        ("<<Total solo mano de obra>>", clp(datos["total_mano_obra"])),
        ("<<Total Rptos>>", clp(datos["total_repuestos"])),
        ("<<Total Repuestos>>", clp(datos["total_repuestos"])),
        ("<<Total Otros>>", clp(datos["total_otros"])),
        ("<<Total>>", clp(datos["total_final"])),
        ("<<Observaciones>>", obs_final),
    ]

    reemplazar_texto_docx(doc, reemplazos)

    nombre = f"Cotizacion_{limpiar_nombre_archivo(numero)}_{limpiar_nombre_archivo(datos['cliente'])}.docx"
    salida = OUTPUT_DIR / nombre
    doc.save(str(salida))
    return salida


def convertir_a_pdf(docx_path: Path) -> Path | None:
    posibles = [shutil.which("soffice"), "/usr/bin/soffice", "/usr/local/bin/soffice"]
    soffice = next((p for p in posibles if p and os.path.exists(p)), None)
    if not soffice:
        return None

    cmd = [
        soffice,
        "--headless",
        "--convert-to",
        "pdf",
        "--outdir",
        str(docx_path.parent),
        str(docx_path),
    ]
    result = subprocess.run(cmd, capture_output=True, text=True, timeout=90)
    if result.returncode != 0:
        return None

    pdf_path = docx_path.with_suffix(".pdf")
    return pdf_path if pdf_path.exists() else None

# =========================================================
# INICIALIZACIÓN
# =========================================================
init_db()

# =========================================================
# CABECERA
# =========================================================
h1, h2 = st.columns([1.1, 5])
with h1:
    if LOGO_HEADER.exists():
        st.image(str(LOGO_HEADER), width=210)
with h2:
    st.markdown("<div class='main-title'>Cotizador de reparaciones</div>", unsafe_allow_html=True)
    st.markdown("<div class='sub-title'>Charles Servicio Automotriz · San Bernardo, Chile</div>", unsafe_allow_html=True)

st.divider()

tab_nueva, tab_historial, tab_dashboard, tab_config = st.tabs([
    "🧾 Nueva cotización",
    "📚 Historial y estados",
    "📊 Dashboard",
    "⚙️ Base / archivos",
])

# =========================================================
# TAB 1 - NUEVA COTIZACIÓN
# =========================================================
with tab_nueva:
    st.subheader("Nueva cotización")
    st.caption("Completa los datos, agrega mano de obra, repuestos y otros cobros. La aplicación calcula el total y guarda todo en base de datos + Excel.")

    with st.form("form_datos_cliente"):
        st.markdown("### 1) Datos del cliente y vehículo")
        c1, c2, c3 = st.columns(3)
        with c1:
            fecha = st.date_input("Fecha", value=date.today())
            cliente = st.text_input("Cliente *")
            contacto = st.text_input("Contacto celular")
            mail = st.text_input("Mail del cliente")
        with c2:
            patente = st.text_input("Patente")
            marca = st.text_input("Marca vehículo")
            modelo = st.text_input("Modelo vehículo")
            anio = st.text_input("Año")
        with c3:
            vin = st.text_input("VIN / Chasis")
            kilometraje = st.text_input("Kilometraje")
            observaciones = st.text_area("Observaciones", height=130)

        guardar_datos = st.form_submit_button("Guardar datos de cabecera", use_container_width=True)

    st.markdown("### 2) Detalle de cobros")
    e1, e2 = st.columns(2)
    with e1:
        df_mo_raw = editor_items("Mano de obra / trabajos a realizar", "editor_mo", "Diagnóstico y revisión general")
    with e2:
        df_rep_raw = editor_items("Repuestos comprometidos", "editor_rep", "Repuesto / material")

    df_otros_raw = editor_items("Otros cobros", "editor_otros", "Traslado / insumo / cargo adicional")

    df_mo = limpiar_df_items(df_mo_raw, "Mano de obra")
    df_rep = limpiar_df_items(df_rep_raw, "Repuestos")
    df_otros = limpiar_df_items(df_otros_raw, "Otros")
    items = pd.concat([df_mo, df_rep, df_otros], ignore_index=True)

    total_mo = float(df_mo["total"].sum()) if not df_mo.empty else 0
    total_rep = float(df_rep["total"].sum()) if not df_rep.empty else 0
    total_otros = float(df_otros["total"].sum()) if not df_otros.empty else 0
    total_final = total_mo + total_rep + total_otros

    st.markdown("### 3) Resumen automático")
    r1, r2, r3, r4 = st.columns(4)
    r1.metric("Total mano de obra", clp(total_mo))
    r2.metric("Total repuestos", clp(total_rep))
    r3.metric("Total otros", clp(total_otros))
    r4.metric("Total final", clp(total_final))

    st.markdown("---")
    generar = st.button("Generar cotización Word / PDF y guardar en base", type="primary", use_container_width=True)

    if generar:
        if not cliente.strip():
            st.error("Debes ingresar el nombre del cliente.")
        elif items.empty:
            st.error("Debes ingresar al menos un ítem de mano de obra, repuestos u otros.")
        elif total_final <= 0:
            st.error("El total final debe ser mayor a cero.")
        else:
            numero = numero_siguiente()
            creado_en = ahora_santiago().strftime("%Y-%m-%d %H:%M:%S")
            datos = {
                "numero_cotizacion": numero,
                "fecha": fecha.isoformat(),
                "fecha_mostrar": fecha.strftime("%d-%m-%Y"),
                "cliente": cliente.strip(),
                "contacto": contacto.strip(),
                "mail": mail.strip(),
                "patente": patente.strip(),
                "marca": marca.strip(),
                "modelo": modelo.strip(),
                "anio": anio.strip(),
                "vin": vin.strip(),
                "kilometraje": kilometraje.strip(),
                "observaciones": observaciones.strip(),
                "total_mano_obra": total_mo,
                "total_repuestos": total_rep,
                "total_otros": total_otros,
                "total_final": total_final,
                "creado_en": creado_en,
            }

            try:
                with st.spinner("Generando documento y guardando registro..."):
                    docx_path = generar_docx_cotizacion(datos, items)
                    pdf_path = convertir_a_pdf(docx_path)
                    datos["docx_path"] = str(docx_path)
                    datos["pdf_path"] = str(pdf_path) if pdf_path else ""
                    cot_id = guardar_cotizacion(datos, items)
                    actualizar_archivos_cotizacion(cot_id, str(docx_path), str(pdf_path) if pdf_path else "")

                st.success(f"Cotización {numero} generada y guardada correctamente.")
                st.write(f"**Cliente:** {cliente}")
                st.write(f"**Total final:** {clp(total_final)}")

                d1, d2 = st.columns(2)
                with d1:
                    with open(docx_path, "rb") as f:
                        st.download_button(
                            "Descargar Word",
                            data=f.read(),
                            file_name=docx_path.name,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True,
                        )
                with d2:
                    if pdf_path and Path(pdf_path).exists():
                        with open(pdf_path, "rb") as f:
                            st.download_button(
                                "Descargar PDF",
                                data=f.read(),
                                file_name=Path(pdf_path).name,
                                mime="application/pdf",
                                use_container_width=True,
                            )
                    else:
                        st.info("PDF no disponible. Para habilitarlo instala LibreOffice/soffice o usa packages.txt en Streamlit Cloud.")

            except Exception as e:
                st.error(f"No fue posible generar la cotización: {e}")

# =========================================================
# TAB 2 - HISTORIAL Y ESTADOS
# =========================================================
with tab_historial:
    st.subheader("Historial y seguimiento de estados")
    df = cargar_cotizaciones()

    if df.empty:
        st.info("Aún no hay cotizaciones guardadas.")
    else:
        f1, f2, f3 = st.columns(3)
        with f1:
            filtro_estado = st.multiselect("Estado", ESTADOS, default=ESTADOS)
        with f2:
            texto_busqueda = st.text_input("Buscar cliente / patente / número")
        with f3:
            orden = st.selectbox("Orden", ["Más recientes", "Más antiguas", "Mayor total", "Menor total"])

        vista = df.copy()
        vista = vista[vista["estado"].isin(filtro_estado)]
        if texto_busqueda.strip():
            q = texto_busqueda.strip().lower()
            vista = vista[
                vista["cliente"].fillna("").str.lower().str.contains(q)
                | vista["patente"].fillna("").str.lower().str.contains(q)
                | vista["numero_cotizacion"].fillna("").str.lower().str.contains(q)
            ]

        if orden == "Más antiguas":
            vista = vista.sort_values("id", ascending=True)
        elif orden == "Mayor total":
            vista = vista.sort_values("total_final", ascending=False)
        elif orden == "Menor total":
            vista = vista.sort_values("total_final", ascending=True)
        else:
            vista = vista.sort_values("id", ascending=False)

        vista_mostrar = vista[[
            "id", "numero_cotizacion", "fecha", "cliente", "contacto", "patente",
            "marca", "modelo", "total_mano_obra", "total_repuestos", "total_otros",
            "total_final", "estado", "motivo_estado", "fecha_estado"
        ]].copy()

        for col in ["total_mano_obra", "total_repuestos", "total_otros", "total_final"]:
            vista_mostrar[col] = vista_mostrar[col].apply(clp)

        st.dataframe(vista_mostrar, use_container_width=True, hide_index=True)

        st.markdown("### Cambiar estado de una cotización")
        opciones = [f"{row.id} | {row.numero_cotizacion} | {row.cliente} | {clp(row.total_final)} | {row.estado}" for row in vista.itertuples()]
        seleccion = st.selectbox("Selecciona cotización", opciones)
        cot_id = int(seleccion.split("|")[0].strip())

        c1, c2 = st.columns([1, 2])
        with c1:
            nuevo_estado = st.selectbox("Nuevo estado", ESTADOS)
        with c2:
            motivo = st.text_input("Comentario / motivo", placeholder="Ej.: Cliente acepta presupuesto / No acepta por precio / En espera de repuestos")

        if st.button("Actualizar estado", use_container_width=True):
            actualizar_estado(cot_id, nuevo_estado, motivo)
            st.success("Estado actualizado y base Excel sincronizada.")
            st.rerun()

        st.markdown("### Detalle de ítems")
        df_items_sel = cargar_items(cot_id)
        if not df_items_sel.empty:
            df_items_show = df_items_sel[["categoria", "descripcion", "cantidad", "valor_unitario", "total"]].copy()
            df_items_show["valor_unitario"] = df_items_show["valor_unitario"].apply(clp)
            df_items_show["total"] = df_items_show["total"].apply(clp)
            st.dataframe(df_items_show, use_container_width=True, hide_index=True)

# =========================================================
# TAB 3 - DASHBOARD
# =========================================================
with tab_dashboard:
    st.subheader("Dashboard de cotizaciones")
    df = cargar_cotizaciones()

    if df.empty:
        st.info("Aún no hay información para graficar.")
    else:
        df["fecha_dt"] = pd.to_datetime(df["fecha"], errors="coerce")
        total_cot = len(df)
        total_monto = df["total_final"].sum()
        total_aceptado = df.loc[df["estado"] == "Aceptada", "total_final"].sum()
        cantidad_aceptada = int((df["estado"] == "Aceptada").sum())
        tasa_aceptacion = (cantidad_aceptada / total_cot * 100) if total_cot else 0

        k1, k2, k3, k4 = st.columns(4)
        k1.metric("Cotizaciones", total_cot)
        k2.metric("Monto cotizado", clp(total_monto))
        k3.metric("Monto aceptado", clp(total_aceptado))
        k4.metric("Tasa aceptación", f"{tasa_aceptacion:.1f}%")

        st.markdown("### Monto por estado")
        resumen_estado = df.groupby("estado", as_index=False).agg(
            cantidad=("id", "count"),
            monto=("total_final", "sum")
        )
        fig_estado = px.bar(resumen_estado, x="estado", y="monto", text="cantidad", title="Monto total por estado")
        fig_estado.update_layout(yaxis_title="Monto CLP", xaxis_title="Estado")
        st.plotly_chart(fig_estado, use_container_width=True)

        st.markdown("### Evolución mensual")
        df_mes = df.dropna(subset=["fecha_dt"]).copy()
        df_mes["mes"] = df_mes["fecha_dt"].dt.to_period("M").astype(str)
        resumen_mes = df_mes.groupby("mes", as_index=False).agg(
            cotizaciones=("id", "count"),
            monto=("total_final", "sum")
        )
        fig_mes = px.line(resumen_mes, x="mes", y="monto", markers=True, title="Monto cotizado por mes")
        fig_mes.update_layout(yaxis_title="Monto CLP", xaxis_title="Mes")
        st.plotly_chart(fig_mes, use_container_width=True)

        st.markdown("### Distribución del total")
        suma_partidas = pd.DataFrame({
            "partida": ["Mano de obra", "Repuestos", "Otros"],
            "monto": [df["total_mano_obra"].sum(), df["total_repuestos"].sum(), df["total_otros"].sum()],
        })
        fig_partidas = px.pie(suma_partidas, names="partida", values="monto", title="Composición general de cotizaciones")
        st.plotly_chart(fig_partidas, use_container_width=True)

# =========================================================
# TAB 4 - BASE / ARCHIVOS
# =========================================================
with tab_config:
    st.subheader("Base de datos y archivos")
    st.write("La aplicación guarda la información en SQLite y mantiene una planilla Excel sincronizada con dos hojas: `cotizaciones` y `detalle_items`.")

    c1, c2, c3 = st.columns(3)
    with c1:
        if EXCEL_FILE.exists():
            with open(EXCEL_FILE, "rb") as f:
                st.download_button(
                    "Descargar base Excel",
                    data=f.read(),
                    file_name=EXCEL_FILE.name,
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True,
                )
        else:
            st.info("La planilla se creará cuando generes la primera cotización.")

    with c2:
        if DB_FILE.exists():
            with open(DB_FILE, "rb") as f:
                st.download_button(
                    "Descargar base SQLite",
                    data=f.read(),
                    file_name=DB_FILE.name,
                    mime="application/octet-stream",
                    use_container_width=True,
                )

    with c3:
        if st.button("Forzar actualización Excel", use_container_width=True):
            sincronizar_excel()
            st.success("Excel actualizado.")

    st.markdown("### Archivos generados")
    archivos = sorted(OUTPUT_DIR.glob("*"), key=lambda p: p.stat().st_mtime, reverse=True)
    if not archivos:
        st.info("Aún no hay documentos generados.")
    else:
        for archivo in archivos[:20]:
            with open(archivo, "rb") as f:
                st.download_button(
                    f"Descargar {archivo.name}",
                    data=f.read(),
                    file_name=archivo.name,
                    use_container_width=True,
                )
